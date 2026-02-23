"""Word 文档加载器"""
from typing import List
from docx import Document as DocxDocument
from src.loaders.base import BaseLoader, Document, CHUNKING_STRATEGY, STRATEGY_REGULAR
from src.chunking.splitter import get_text_splitter
from src.logger import get_logger

logger = get_logger("docx_loader")


class DocxLoader(BaseLoader):
    """Word (.docx) 文档加载器 - 使用常规切分器"""

    def load(self, path: str) -> List[Document]:
        """
        加载 Word 文档

        提取所有段落内容，然后使用常规切分器切分。

        Args:
            path: Word 文件路径

        Returns:
            切分后的文档列表
        """
        path_obj = self.validate_file_path(path, file_type="Word")

        doc = DocxDocument(path)
        paragraphs = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)

        content = "\n".join(paragraphs)

        # 使用常规切分器切分
        chunked_docs = []
        text_splitter = get_text_splitter()
        chunks = text_splitter.split_text(content)

        for i, chunk in enumerate(chunks):
            chunked_doc = Document(
                content=chunk,
                metadata={
                    "type": "docx",
                    "paragraphs_count": len(doc.paragraphs),
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    CHUNKING_STRATEGY: STRATEGY_REGULAR,
                },
                source=str(path_obj),
            )
            chunked_docs.append(chunked_doc)

        logger.info(f"DOCX 切分完成: {len(chunked_docs)} 个块 (来自 {len(doc.paragraphs)} 段落)")
        return chunked_docs
